"""
ingest.py — Jalankan sekali di lokal untuk load, chunk, embed, dan simpan ke ChromaDB.
Output: chroma_db/ (di-commit ke GitHub) + bm25_chunks.json (untuk BM25Retriever di retriever.py)
"""

import os
import json
import time
import re
from pathlib import Path

from dotenv import load_dotenv
import docx
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from tenacity import (
    retry,
    retry_if_exception_type,
    stop_after_attempt,
    wait_exponential,
)

load_dotenv()

# ─── Konfigurasi ──────────────────────────────────────────────────────────────

DOCS_DIR = Path("docs")
CHROMA_DIR = "chroma_db"
BM25_CHUNKS_FILE = "bm25_chunks.json"
EMBED_MODEL = "gemini-embedding-001"
BATCH_SIZE = 5    # chunk per batch saat embed
BATCH_DELAY = 3.0  # detik antar batch (hindari rate limit)

# Metadata untuk 9 case report (document-level chunking)
CASE_REPORT_META: dict[str, dict] = {
    "case_01_canter_cold_start_noise_timing_chain.docx":
        {"doc_type": "case_report", "vehicle_model": "Canter",       "chunk_id": "case_01"},
    "case_02_canter_hard_starting_udara_di_bbm.docx":
        {"doc_type": "case_report", "vehicle_model": "Canter",       "chunk_id": "case_02"},
    "case_03_canter_brake_fade_turunan_panjang.docx":
        {"doc_type": "case_report", "vehicle_model": "Canter",       "chunk_id": "case_03"},
    "case_04_fighter_overheat_oring_thermostat_housing.docx":
        {"doc_type": "case_report", "vehicle_model": "Fighter X",    "chunk_id": "case_04"},
    "case_05_fighter_power_loss_filter_bbm_sekunder.docx":
        {"doc_type": "case_report", "vehicle_model": "Fighter X",    "chunk_id": "case_05"},
    "case_06_fighter_vibrasi_balance_weight_propshaft.docx":
        {"doc_type": "case_report", "vehicle_model": "Fighter X",    "chunk_id": "case_06"},
    "case_07_fighter_asap_putih_air_biasa_coolant.docx":
        {"doc_type": "case_report", "vehicle_model": "Fighter X",    "chunk_id": "case_07"},
    "case_08_umum_keausan_dini_intake_hose_retak.docx":
        {"doc_type": "case_report", "vehicle_model": "Lintas model", "chunk_id": "case_08"},
    "case_09_umum_mesin_mati_solar_terkontaminasi_air.docx":
        {"doc_type": "case_report", "vehicle_model": "Lintas model", "chunk_id": "case_09"},
}

# Metadata untuk 4 SOP/troubleshooting (section-based chunking)
SOP_META: dict[str, dict] = {
    "sop_canter_fe74hd.docx":
        {"doc_type": "sop",             "vehicle_model": "Canter"},
    "sop_fighter_x_fn61fsl.docx":
        {"doc_type": "sop",             "vehicle_model": "Fighter X"},
    "troubleshooting_canter_fe74hd.docx":
        {"doc_type": "troubleshooting", "vehicle_model": "Canter"},
    "troubleshooting_fighter_x_fn61fsl.docx":
        {"doc_type": "troubleshooting", "vehicle_model": "Fighter X"},
}


# ─── Chunking ─────────────────────────────────────────────────────────────────

def chunk_case_report(filepath: Path, meta: dict) -> list[Document]:
    """Satu file = satu chunk utuh (document-level chunking)."""
    doc = docx.Document(filepath)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return [Document(
        page_content=full_text,
        metadata={"source": filepath.name, **meta},
    )]


import re  # tambahkan di bagian import atas file

def chunk_sop_by_section(filepath: Path, base_meta: dict) -> list[Document]:
    """Pecah per section — deteksi via bold + diawali angka (1., 1.1, 4.2, dst)."""
    doc = docx.Document(filepath)
    chunks: list[Document] = []
    current_heading: str | None = None
    buffer: list[str] = []
    section_idx = 1

    def is_section_header(para) -> bool:
        text = para.text.strip()
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        starts_with_number = bool(re.match(r'^\d+', text))
        return is_bold and starts_with_number

    def flush():
        nonlocal section_idx
        if not buffer:
            return
        heading_prefix = f"{current_heading}\n" if current_heading else ""
        text = heading_prefix + "\n".join(buffer)
        chunk_id = f"{base_meta['doc_type']}_{filepath.stem}_sec{section_idx}"
        chunks.append(Document(
            page_content=text,
            metadata={
                "source": filepath.name,
                "section": current_heading or "intro",
                "chunk_id": chunk_id,
                **base_meta,
            },
        ))
        section_idx += 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if is_section_header(para):
            flush()
            current_heading = text
            buffer = []
        else:
            buffer.append(text)

    flush()
    return chunks


# ─── Load semua dokumen ───────────────────────────────────────────────────────

def load_all_chunks() -> list[Document]:
    all_chunks: list[Document] = []

    print("=== Case reports (document-level chunking) ===")
    for filename, meta in CASE_REPORT_META.items():
        path = DOCS_DIR / filename
        if not path.exists():
            print(f"  [SKIP] {filename} — tidak ditemukan")
            continue
        chunks = chunk_case_report(path, meta)
        print(f"  {filename} → {len(chunks)} chunk")
        all_chunks.extend(chunks)

    print("\n=== SOP & Troubleshooting (section-based chunking) ===")
    for filename, meta in SOP_META.items():
        path = DOCS_DIR / filename
        if not path.exists():
            print(f"  [SKIP] {filename} — tidak ditemukan")
            continue
        chunks = chunk_sop_by_section(path, meta)
        print(f"  {filename} → {len(chunks)} chunk")
        all_chunks.extend(chunks)

    print(f"\nTotal: {len(all_chunks)} chunks")
    return all_chunks


# ─── Simpan chunks untuk BM25 ────────────────────────────────────────────────

def save_bm25_chunks(all_chunks: list[Document]) -> None:
    """Simpan sebagai JSON agar retriever.py bisa recreate Document objects tanpa re-ingest."""
    data = [
        {"page_content": doc.page_content, "metadata": doc.metadata}
        for doc in all_chunks
    ]
    with open(BM25_CHUNKS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"BM25 chunks disimpan → {BM25_CHUNKS_FILE} ({len(data)} entries)")


# ─── Embed + simpan ke ChromaDB ──────────────────────────────────────────────

class _RateLimitError(Exception):
    pass


def _is_rate_limit(exc: Exception) -> bool:
    msg = str(exc).lower()
    return any(kw in msg for kw in ("429", "quota", "rate limit", "resource exhausted"))


@retry(
    retry=retry_if_exception_type(_RateLimitError),
    wait=wait_exponential(multiplier=10, min=10, max=120),
    stop=stop_after_attempt(6),
    reraise=True,
)
def _add_batch(vectorstore: Chroma, batch: list[Document]) -> None:
    try:
        vectorstore.add_documents(batch)
    except Exception as exc:
        if _is_rate_limit(exc):
            print(f"  Rate limit — akan retry otomatis...")
            raise _RateLimitError(str(exc)) from exc
        raise


def build_vectorstore(all_chunks: list[Document], embeddings: GoogleGenerativeAIEmbeddings) -> Chroma:
    print("\n=== Membangun ChromaDB ===")
    total_batches = (len(all_chunks) + BATCH_SIZE - 1) // BATCH_SIZE

    # Inisialisasi ChromaDB dari batch pertama
    first_batch = all_chunks[:BATCH_SIZE]
    print(f"Batch 1/{total_batches} — inisialisasi ({len(first_batch)} chunks)...")
    vectorstore = Chroma.from_documents(
        documents=first_batch,
        embedding=embeddings,
        persist_directory=CHROMA_DIR,
    )

    # Tambah sisa chunks dalam batch
    remaining = all_chunks[BATCH_SIZE:]
    for i in range(0, len(remaining), BATCH_SIZE):
        batch = remaining[i : i + BATCH_SIZE]
        batch_num = (i // BATCH_SIZE) + 2
        print(f"Batch {batch_num}/{total_batches} ({len(batch)} chunks)...")
        time.sleep(BATCH_DELAY)
        _add_batch(vectorstore, batch)

    return vectorstore


# ─── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError(
            "GOOGLE_API_KEY tidak ditemukan. "
            "Tambahkan ke file .env atau set sebagai environment variable."
        )

    all_chunks = load_all_chunks()
    if not all_chunks:
        raise RuntimeError("Tidak ada chunk yang berhasil di-load. Cek folder docs/.")

    print("\n=== Menyimpan BM25 chunks ===")
    save_bm25_chunks(all_chunks)

    print("\n=== Inisialisasi embedding model ===")
    print(f"Model: {EMBED_MODEL}")
    embeddings = GoogleGenerativeAIEmbeddings(
        model=EMBED_MODEL,
        google_api_key=api_key,
    )

    vectorstore = build_vectorstore(all_chunks, embeddings)

    try:
        total_vectors = vectorstore._collection.count()
        print(f"\n=== Selesai ===")
        print(f"ChromaDB: '{CHROMA_DIR}/' ({total_vectors} vectors)")
    except Exception:
        print(f"\n=== Selesai ===")
        print(f"ChromaDB disimpan di '{CHROMA_DIR}/'")

    print(f"BM25 chunks: '{BM25_CHUNKS_FILE}'")
    print("\nLangkah selanjutnya:")
    print("  git add chroma_db/ bm25_chunks.json")
    print("  git commit -m 'add: ChromaDB dan BM25 chunks hasil ingest'")


if __name__ == "__main__":
    main()
